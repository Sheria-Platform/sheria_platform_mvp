# pipelines/ingestion/loaders/docx.py
import docx
import io

def parse_docx_bytes(file_bytes: bytes, filename: str):
    """Parses .docx files extracting text and simple tables."""
    doc = docx.Document(io.BytesIO(file_bytes))
    full_text = []
    
    for para in doc.paragraphs:
        if para.text.strip():
            full_text.append(para.text)
            
    for table in doc.tables:
        for row in table.rows:
            row_text = "\t".join(cell.text.strip() for cell in row.cells)
            full_text.append(row_text)
            
    return "\n\n".join(full_text), {"filename": filename, "type": "docx"}


if __name__ == "__main__":
    # Example usage
    with open("scalable-rag-core/pipelines/ingestion/tests/sample.docx", "rb") as f:
        file_bytes = f.read()
    
    text, metadata = parse_docx_bytes(file_bytes, "example.docx")
    print("Extracted Text:\n", text)
    print("Metadata:\n", metadata)